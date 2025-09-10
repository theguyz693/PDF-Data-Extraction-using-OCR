import os
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
import easyocr
from paddleocr import PaddleOCR
from pdfminer.high_level import extract_text as pdfminer_extract_text

# Configure Tesseract manually if needed on Windows
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Initialize OCR engines
easyocr_reader = easyocr.Reader(['en'], gpu=False)
paddleocr_reader = PaddleOCR(use_angle_cls=True, lang='en')

class PDFExtractor:
    def __init__(self, pdf_path, engines=None, tools=None):
        self.pdf_path = pdf_path
        self.engines = [e.strip().lower() for e in (engines or ["tesseract"])]
        self.tools = [t.strip().lower() for t in (tools or ["pdfplumber"])]

    # -------- PDF Parsing Tools -------- #
    def _extract_with_plumber(self, page):
        return page.extract_text(x_tolerance=1, y_tolerance=1, layout=True) or ""

    def _extract_with_pymupdf(self, doc, page_num):
        return doc.load_page(page_num).get_text("text")

    def _extract_with_pdfminer(self):
        try:
            return pdfminer_extract_text(self.pdf_path)
        except Exception as e:
            print(f"PDFMiner error: {e}")
            return ""

    # -------- OCR Engines -------- #
    def _ocr_tesseract(self, image):
        return pytesseract.image_to_string(image, lang="eng")

    def _ocr_easyocr(self, image_path):
        results = easyocr_reader.readtext(image_path)
        return " ".join([res[1] for res in results])

    def _ocr_paddleocr(self, image_path):
        results = paddleocr_reader.ocr(image_path, cls=True)
        return " ".join([line[1][0] for line in results[0]]) if results else ""

    # -------- Combined Extraction -------- #
    def extract(self):
        extracted_data = {}
        try:
            fitz_doc = fitz.open(self.pdf_path)

            with pdfplumber.open(self.pdf_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = ""

                    # Try all selected tools
                    for tool in self.tools:
                        if tool == "pdfplumber":
                            text = self._extract_with_plumber(page)
                        elif tool == "pymupdf":
                            text = self._extract_with_pymupdf(fitz_doc, i)
                        elif tool == "pdfminer":
                            if i == 0:  # run once for whole doc
                                text = self._extract_with_pdfminer()
                        if text:
                            break

                    # Fallback: OCR
                    if not text:
                        pil_image = page.to_image(resolution=300).original
                        temp_path = f"temp_page_{i}.png"
                        pil_image.save(temp_path)

                        for engine in self.engines:
                            if engine == "tesseract":
                                text = self._ocr_tesseract(pil_image)
                            elif engine == "easyocr":
                                text = self._ocr_easyocr(temp_path)
                            elif engine == "paddleocr":
                                text = self._ocr_paddleocr(temp_path)
                            if text:
                                break
                        os.remove(temp_path)

                    extracted_data[i+1] = text
            return extracted_data

        except Exception as e:
            print(f"Error extracting: {e}")
            return {}

    # -------- Save Output -------- #
    def save_to_txt(self, data, output_path):
        with open(output_path, "w", encoding="utf-8") as f:
            for page, text in data.items():
                f.write(f"--- Page {page} ---\n{text}\n\n")
        print(f"✅ Output saved to {output_path}")

    def save_to_docx(self, data, output_path):
        doc = Document()
        for page, text in data.items():
            doc.add_heading(f"Page {page}", level=1)
            doc.add_paragraph(text)
        doc.save(output_path)
        print(f"✅ Output saved to {output_path}")

# -------- CLI Entry -------- #
def main():
    pdf_path = input("📄 Enter PDF path: ").strip()
    if not os.path.exists(pdf_path):
        print("❌ File not found.")
        return

    engines = input("⚙️  Choose OCR engines (comma sep: tesseract,easyocr,paddleocr): ").strip().split(",")
    tools = input("🔧 Choose tools (comma sep: pdfplumber,pymupdf,pdfminer): ").strip().split(",")
    output_format = input("💾 Choose output format (txt/docx): ").strip().lower()

    # File naming
    custom_name = input("✍️  Enter output filename (without extension) or press Enter for default: ").strip()
    base_name = custom_name if custom_name else os.path.splitext(pdf_path)[0] + "_output"

    extractor = PDFExtractor(pdf_path, engines, tools)
    data = extractor.extract()
    if not data:
        print("❌ No data extracted.")
        return

    if output_format == "txt":
        extractor.save_to_txt(data, f"{base_name}.txt")
    elif output_format == "docx":
        extractor.save_to_docx(data, f"{base_name}.docx")
    else:
        print("❌ Invalid format. Choose txt or docx.")

if __name__ == "__main__":
    main()
